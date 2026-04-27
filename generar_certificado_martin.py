from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_certificate_martin():
    doc = Document()

    # Título Principal
    title = doc.add_heading('CERTIFICADO DE GARANTÍA Y MANUAL DE CUIDADOS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph('Lithium Electrónica – Tecnología de Litio Fase Industrial')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True

    doc.add_paragraph('\n')

    # Datos del Equipo
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    data = [
        ('CLIENTE:', 'Martin Santos'),
        ('PRODUCTO:', 'Batería de Litio Pro (24V 30Ah)'),
        ('CONFIGURACIÓN:', '7S6P - Celdas EVE 21700 / BMS Daly High-Current'),
        ('FECHA ENTREGA:', '27 / 04 / 2026'),
        ('VENCIMIENTO GARANTÍA:', '27 / 04 / 2027 (12 Meses)')
    ]

    for i, (key, value) in enumerate(data):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph('\n')

    # Especificaciones de Carga
    h_carga = doc.add_heading('ESPECIFICACIONES DEL CARGADOR', level=1)
    p_carga = doc.add_paragraph()
    p_carga.add_run('TIPO DE CARGADOR: ').bold = True
    p_carga.add_run('Cargador específico para Li-ion (7S).\n')
    p_carga.add_run('VOLTAJE DE CARGA: ').bold = True
    p_carga.add_run('29.4V DC.\n')
    p_carga.add_run('IMPORTANTE: ').bold = True
    p_carga.add_run('Queda terminantemente prohibido el uso de cargadores para baterías de Plomo/Gel o cargadores de coche. Su uso destruye la química del litio y anula la garantía.')

    # Términos de la Garantía
    h_garantia = doc.add_heading('COBERTURA Y ANULACIÓN', level=1)
    
    terms = [
        "COBERTURA: La garantía cubre defectos de fabricación en el pack de celdas y falla electrónica del BMS.",
        "ANULACIÓN POR HUMEDAD: El ingreso de agua o humedad interna anula la garantía. No lavar a presión ni usar bajo lluvia intensa sin protección adicional.",
        "ANULACIÓN POR CALOR: Exponer la batería a temperaturas superiores a 50°C (sol directo prolongado) daña las celdas y anula la cobertura.",
        "MAL USO: Golpes mecánicos, apertura de los sellos de seguridad o descarga profunda por debajo de 19V anulan la garantía automáticamente."
    ]

    for term in terms:
        doc.add_paragraph(term, style='List Bullet')

    # Guía de Cuidado (Vida Útil)
    h_uso = doc.add_heading('GUÍA DE USO Y VIDA ÚTIL', level=1)
    
    advices = [
        "CARGA SALUDABLE: Se recomienda cargar la batería cuando llegue al 20-30% de su capacidad.",
        "ORDEN DE CONEXIÓN: Conecte primero el cargador a la pared y luego a la batería.",
        "EQUILIBRIO DE CELDAS: Una vez al mes, deje la batería conectada 2 horas extra después de que la luz del cargador pase a verde.",
        "INTEMPERIE: Si el vehículo duerme afuera, se recomienda retirar la batería o cubrirla con funda impermeable."
    ]

    for advice in advices:
        doc.add_paragraph(advice, style='List Number')

    doc.add_paragraph('\n')
    
    # Firmas
    f_table = doc.add_table(rows=1, cols=2)
    f_table.rows[0].cells[0].text = "__________________________\nFirma del Técnico\nLithium Electrónica"
    f_table.rows[0].cells[1].text = "__________________________\nFirma del Cliente\nMartin Santos"
    
    doc.save('Certificado_Garantia_Martin_Santos.docx')
    print("Certificado generado con éxito.")

if __name__ == "__main__":
    create_certificate_martin()
