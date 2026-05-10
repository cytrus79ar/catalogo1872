import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_manual():
    doc = Document()
    
    # Header with Logo and Date
    header_table = doc.add_table(rows=1, cols=2)
    header_table.width = Inches(6)
    
    # Logo
    logo_path = r'c:\Users\Marcos\Documents\LithiumBateriasPro\media\banner mercado.jpg'
    if os.path.exists(logo_path):
        run = header_table.cell(0, 0).paragraphs[0].add_run()
        run.add_picture(logo_path, width=Inches(2))
    
    # Date and Time
    now = datetime.now()
    dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
    p_dt = header_table.cell(0, 1).paragraphs[0]
    p_dt.text = f"Generado el: {dt_string}"
    p_dt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_heading('MANUAL DE INSTRUCCIONES DE USO', 0)
    
    doc.add_paragraph('Has adquirido una batería de litio de fase industrial. Para asegurar su rendimiento óptimo y máxima vida útil, seguí estas indicaciones:')
    
    doc.add_heading('1. Protocolo de Carga', level=1)
    doc.add_paragraph('Utilizá únicamente el cargador provisto o recomendado por Lithium Electrónica.', style='List Bullet')
    doc.add_paragraph('Conectá primero el cargador a la batería y luego al tomacorriente de 220V.', style='List Bullet')
    doc.add_paragraph('Es recomendable cargar la batería después de cada uso, sin esperar a que se agote por completo.', style='List Bullet')
    doc.add_paragraph('IMPORTANTE: Esperá 15-20 minutos después de usar el vehículo antes de ponerlo a cargar.', style='List Bullet')
    
    doc.add_heading('2. Cuidado y Mantenimiento', level=1)
    doc.add_paragraph('Evitá temperaturas extremas (más de 50°C). No dejes el vehículo al sol prolongadamente.', style='List Bullet')
    doc.add_paragraph('Mantené el pack alejado de la humedad excesiva. No es sumergible.', style='List Bullet')
    doc.add_paragraph('Si no vas a usar la batería por más de un mes, guardala con un 60% de carga.', style='List Bullet')
    
    doc.add_heading('3. Seguridad', level=1)
    doc.add_paragraph('No intentes abrir, perforar ni modificar el pack. Contiene alta densidad energética.', style='List Bullet')
    
    doc.add_paragraph('\nLithium Electrónica - San Rafael, Mendoza', style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(r'c:\Users\Marcos\Documents\LithiumBateriasPro\01_Ventas_y_Marketing\Manual_Instrucciones_Lithium.docx')

def create_garantia():
    doc = Document()
    
    # Header
    header_table = doc.add_table(rows=1, cols=2)
    logo_path = r'c:\Users\Marcos\Documents\LithiumBateriasPro\media\banner mercado.jpg'
    if os.path.exists(logo_path):
        run = header_table.cell(0, 0).paragraphs[0].add_run()
        run.add_picture(logo_path, width=Inches(2))
    
    now = datetime.now()
    dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
    p_dt = header_table.cell(0, 1).paragraphs[0]
    p_dt.text = f"Fecha de Emisión: {dt_string}"
    p_dt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_heading('CERTIFICADO DE GARANTÍA ESCRITA', 0)
    
    doc.add_paragraph('Este documento certifica que el pack de batería adjunto ha sido fabricado bajo estándares industriales y cuenta con el respaldo técnico de Lithium Electrónica.')
    
    # Table for data
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Titular:'
    table.cell(1, 0).text = 'Modelo de Pack:'
    table.cell(2, 0).text = 'N° de Serie / ID:'
    
    doc.add_heading('Cobertura (12 Meses)', level=1)
    doc.add_paragraph('Fallas internas en las celdas de litio (Grado A).', style='List Bullet')
    doc.add_paragraph('Defectos en el sistema electrónico de gestión (BMS).', style='List Bullet')
    doc.add_paragraph('Integridad de conexiones y soldaduras industriales.', style='List Bullet')
    
    doc.add_heading('Exclusiones', level=1)
    doc.add_paragraph('Daños por accidentes, golpes o caídas.', style='List Bullet')
    doc.add_paragraph('Ingreso de agua o líquidos (no sumergible).', style='List Bullet')
    doc.add_paragraph('Uso de cargadores no compatibles.', style='List Bullet')
    doc.add_paragraph('Violación de sellos de seguridad.', style='List Bullet')
    
    doc.add_paragraph('\n\n__________________________          __________________________')
    doc.add_paragraph('      Firma del Cliente                     Sello y Firma Lithium')
    
    doc.save(r'c:\Users\Marcos\Documents\LithiumBateriasPro\01_Ventas_y_Marketing\Certificado_Garantia_Lithium.docx')

if __name__ == "__main__":
    create_manual()
    create_garantia()
    print("Documentos generados exitosamente.")
