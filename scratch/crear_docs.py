import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_certificado():
    doc = docx.Document()
    
    # Header / Logo placeholder
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = "[ LITHIUM ELECTRÓNICA ]"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.bold = True
    hp.runs[0].font.color.rgb = RGBColor(0, 86, 179)
    
    # Title
    doc.add_heading('CERTIFICADO DE GARANTÍA', 0)
    
    p = doc.add_paragraph()
    p.add_run('Titular: ').bold = True
    p.add_run('José García Alonso\n')
    p.add_run('Producto: ').bold = True
    p.add_run('Batería 36V 10Ah – Celdas EVE 21700/50e\n')
    p.add_run('Número de Serie: ').bold = True
    p.add_run('LE-36-10-EVE-2026-00123\n')
    p.add_run('Fecha de Compra: ').bold = True
    p.add_run('07/05/2026\n')
    p.add_run('Fecha de Vencimiento: ').bold = True
    p.add_run('07/05/2027 (1 Año de Cobertura)\n')
    
    doc.add_heading('Términos y Condiciones de la Garantía', level=1)
    
    terms = [
        "Cobertura total por defectos de fabricación y materiales.",
        "La garantía es intransferible y válida sólo para el titular indicado en este documento.",
        "Para hacer efectivo cualquier reclamo, es indispensable presentar este certificado junto con el número de serie legible en la batería.",
        "El tiempo de revisión técnico es de 48 a 72 horas hábiles desde la recepción de la batería en el taller."
    ]
    for term in terms:
        doc.add_paragraph(term, style='List Bullet')
        
    doc.add_heading('Causales de Anulación de Garantía', level=1)
    
    anulaciones = [
        "Daños por mal uso, negligencia o accidentes.",
        "Exposición a condiciones extremas: temperaturas superiores a 60°C o luz solar directa prolongada.",
        "Exposición a agua, lluvia o niveles altos de humedad (la batería no es sumergible).",
        "Signos de cortocircuitos o alteraciones en los conectores.",
        "Manipulación no autorizada, apertura de la carcasa o reparaciones externas.",
        "Signos de uso inadecuado reportados por el BMS: sobrecarga sostenida, descargas profundas por debajo del límite de seguridad."
    ]
    for anulación in anulaciones:
        doc.add_paragraph(anulación, style='List Bullet')
        
    doc.add_paragraph('\n')
    p_sign = doc.add_paragraph()
    p_sign.add_run('___________________________________\n')
    p_sign.add_run('Firma Autorizada\nLithium Electrónica - San Rafael, Mendoza')
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save('c:/Users/Marcos/Documents/LithiumBateriasPro/04_Documentacion_y_Legales/certificado_garantia_Jose_Garcia_Alonso.docx')

def create_manual():
    doc = docx.Document()
    
    # Header / Logo placeholder
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = "[ LITHIUM ELECTRÓNICA ]"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.bold = True
    hp.runs[0].font.color.rgb = RGBColor(0, 86, 179)
    
    # Title
    doc.add_heading('MANUAL DE USO Y CUIDADOS', 0)
    
    doc.add_heading('Especificaciones Técnicas', level=1)
    p = doc.add_paragraph()
    p.add_run('Modelo: ').bold = True
    p.add_run('Batería 36V 10Ah\n')
    p.add_run('Celdas: ').bold = True
    p.add_run('EVE 21700/50e (Configuración 10S 2P)\n')
    p.add_run('Voltaje Nominal: ').bold = True
    p.add_run('36V\n')
    p.add_run('Capacidad: ').bold = True
    p.add_run('10Ah (Aprox. 360 Wh)\n')
    p.add_run('Peso: ').bold = True
    p.add_run('Aprox. 2 kg\n')
    p.add_run('BMS Integrado: ').bold = True
    p.add_run('Sistema de gestión inteligente con protección contra sobrecarga, sobredescarga, cortocircuitos y temperatura extrema.')

    doc.add_heading('Instalación y Conexión', level=1)
    instrucciones = [
        "Verifique que el voltaje del controlador y cargador coincidan con los de la batería.",
        "Conecte firmemente los terminales, asegurándose de respetar la polaridad (Rojo = Positivo, Negro = Negativo).",
        "Evite forzar los cables o dejarlos tensos, ya que esto podría causar fatiga en las soldaduras o cortocircuitos."
    ]
    for inst in instrucciones:
        doc.add_paragraph(inst, style='List Bullet')

    doc.add_heading('Carga de la Batería', level=1)
    cargas = [
        "Utilice EXCLUSIVAMENTE cargadores diseñados para baterías de iones de litio (42V).",
        "La corriente de carga recomendada es de 2A a 5A.",
        "Cargue la batería preferentemente en ambientes frescos y ventilados, alejado de materiales inflamables.",
        "Desconecte el cargador una vez que el ciclo haya finalizado o la luz verde se encienda."
    ]
    for c in cargas:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_heading('Uso, Cuidados y Mantenimiento', level=1)
    cuidados = [
        "Protección Solar y Térmica: Nunca deje la batería bajo el sol directo o en vehículos cerrados en días calurosos. Operar la batería a temperaturas superiores a 60°C dañará permanentemente las celdas.",
        "Humedad y Lluvia: La batería NO es impermeable ni sumergible. Proteja la unidad de la lluvia constante, charcos y humedad excesiva. La filtración de agua anula la garantía de manera inmediata.",
        "Cortocircuitos: Nunca puentee los terminales ni utilice objetos metálicos cerca de los conectores expuestos.",
        "Almacenamiento Prolongado: Si no utilizará la batería por más de un mes, asegúrese de guardarla en un ambiente seco (15°C a 25°C) con una carga aproximada del 50%. Realice una carga de mantenimiento cada 3 meses."
    ]
    for c in cuidados:
        doc.add_paragraph(c, style='List Number')
        
    doc.add_heading('Seguridad y Advertencias', level=1)
    seguridad = [
        "Si detecta hinchazón en el encapsulado, olores inusuales o calentamiento excesivo sin estar en uso, desconéctela de inmediato y póngase en contacto con el soporte técnico.",
        "No intente desarmar la batería bajo ninguna circunstancia. El interior contiene energía almacenada y materiales químicos que requieren manipulación profesional.",
        "En caso de siniestro, emplee únicamente extintores apropiados para equipos eléctricos (CO2 o Polvo Químico Seco ABCD)."
    ]
    for s in seguridad:
        doc.add_paragraph(s, style='List Bullet')

    doc.save('c:/Users/Marcos/Documents/LithiumBateriasPro/04_Documentacion_y_Legales/manual_usuario_Jose_Garcia_Alonso.docx')

if __name__ == '__main__':
    create_certificado()
    create_manual()
