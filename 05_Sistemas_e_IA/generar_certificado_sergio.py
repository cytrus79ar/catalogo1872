from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_certificate():
    doc = Document()

    # Título Principal
    title = doc.add_heading('CERTIFICADO DE GARANTÍA TÉCNICA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph('Lithium Electrónica – Especialistas en Movilidad Eléctrica')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True

    doc.add_paragraph('\n')

    # Datos del Cliente
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    data = [
        ('CLIENTE:', 'Sergio Rana'),
        ('PRODUCTO:', 'Batería de Litio High-Performance (36V 10Ah)'),
        ('COMPONENTES:', 'Celdas EVE 21700 / BMS JBD Smart Bluetooth'),
        ('FECHA ENTREGA:', '_____ / _____ / 2026'),
        ('VENCIMIENTO:', '_____ / _____ / 2027 (12 Meses)')
    ]

    for i, (key, value) in enumerate(data):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph('\n')

    # Términos y Condiciones
    h = doc.add_heading('TÉRMINOS Y CONDICIONES', level=1)
    
    terms = [
        "COBERTURA: Defectos de ensamblado, fallas en el BMS y celdas de litio.",
        "SISTEMA SMART: El monitoreo vía App Xiaoxiang es informativo.",
        "ADVERTENCIA: Cualquier modificación de parámetros en la App anula la garantía.",
        "CADUCIDAD: Apertura de sellos, humedad, maltrato mecánico o descarga profunda."
    ]

    for term in terms:
        p = doc.add_paragraph(term, style='List Bullet')

    doc.add_paragraph('\n')

    # Guía de Uso
    h2 = doc.add_heading('GUÍA DE MANTENIMIENTO Y USO RECOMENDADO', level=1)
    
    advices = [
        "CONEXIÓN: Conecte primero a 220V y luego a la batería para evitar chispazos.",
        "TEMPERATURA: Reposo de 30 min tras el uso antes de cargar.",
        "CÁRGALO AL 30%: Nunca deje la batería llegar al 0%.",
        "BALANCEO: Deje conectada 2 horas extra una vez por quincena (Luz Verde).",
        "ALMACENAMIENTO: Guarde al 60-70% si no la usa por más de 15 días."
    ]

    for advice in advices:
        doc.add_paragraph(advice, style='List Number')

    doc.add_paragraph('\n')
    
    # Firmas
    f_table = doc.add_table(rows=1, cols=2)
    f_table.rows[0].cells[0].text = "__________________________\nFirma del Técnico\nMarcos Canales"
    f_table.rows[0].cells[1].text = "__________________________\nFirma del Cliente\nSergio Rana"
    
    doc.save('Certificado_Garantia_Sergio_Rana.docx')

if __name__ == "__main__":
    create_certificate()
